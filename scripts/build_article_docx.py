# MAST: сборка статьи docx из markdown.
# Парсит article/mast_article.md и генерирует article/mast_article.docx
# с заголовками, абзацами, таблицами и встроенными рисунками.

import os
import re
import sys

import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "src"))

from topostress.config import ARTICLE_DIR, FIGURES_DIR

MD = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ARTICLE_DIR, "mast_article.md")
OUT = sys.argv[2] if len(sys.argv) > 2 else os.path.join(ARTICLE_DIR, "mast_article.docx")
if not os.path.isabs(MD):
    MD = os.path.join(ARTICLE_DIR, os.path.basename(MD))

BOLD = re.compile(r"\*\*(.+?)\*\*")
ITAL = re.compile(r"\*(.+?)\*")
FIG = re.compile(r"\((figures/[^)]+\.png)\)")


def add_runs(par, text):
    # Сначала находим полужирные фрагменты, затем курсивные, исключая совпадения,
    # целиком попадающие внутрь **...** (иначе ITAL-регэксп ломает **bold**).
    bolds = list(BOLD.finditer(text))
    bold_spans = [m.span() for m in bolds]
    itals = [m for m in ITAL.finditer(text)
             if not any(bs <= m.start() and m.end() <= be for bs, be in bold_spans)]
    events = ([(m.start(), m.end(), m.group(1), True) for m in bolds]
              + [(m.start(), m.end(), m.group(1), False) for m in itals])
    pos = 0
    for start, end, group, is_bold in sorted(events, key=lambda e: e[0]):
        if start < pos:
            continue
        if start > pos:
            par.add_run(text[pos:start])
        run = par.add_run(group)
        if is_bold:
            run.bold = True
        else:
            run.italic = True
        pos = end
    if pos < len(text):
        par.add_run(text[pos:])


def main():
    lines = open(MD, encoding="utf-8").read().splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # таблица: строка-разделитель |-...-|
        if line.lstrip().startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            data = []
            for row in rows:
                cells = [c.strip() for c in row.strip("|").split("|")]
                if set("".join(cells)) <= set("-: "):
                    continue
                data.append(cells)
            if data:
                table = doc.add_table(rows=len(data), cols=len(data[0]))
                table.style = "Light Grid Accent 1"
                for r, row in enumerate(data):
                    for c, cell in enumerate(row):
                        cell_text = re.sub(r"\*\*", "", cell)
                        table.cell(r, c).text = cell_text
                        if r == 0:
                            for run in table.cell(r, c).paragraphs[0].runs:
                                run.bold = True
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            i += 1
            continue

        # рисунки: строка вида "- Рис. N (figures/x.png) — описание"
        m = FIG.search(line)
        if m and line.lstrip().startswith("-"):
            p = doc.add_paragraph()
            add_runs(p, line.lstrip()[1:].strip())
            fig_path = os.path.join(FIGURES_DIR, os.path.basename(m.group(1)))
            if os.path.exists(fig_path):
                pic = doc.add_picture(fig_path, width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.lstrip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line.lstrip()[2:].strip())
            i += 1
            continue
        if re.match(r"^\d+\.\s", line.strip()):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s", "", line.strip()))
            i += 1
            continue

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, line.strip())
        i += 1

    doc.save(OUT)
    print(f"Сохранено: {OUT}")


if __name__ == "__main__":
    main()
